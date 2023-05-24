import random
import pandas as pd
import docx
import os
import csv
from datetime import datetime
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

INITIAL_WIDTH = 20
INITIAL_HEIGHT = 20
WORD_LIMIT = 15
HEIGHT_LIMIT = 14
WIDTH_LIMIT = 20

words = []
board = []
init = False

def import_words():
    global words
    local_words = []
    with open("wyrazy.csv", encoding="utf8") as csvfile:
        csvreader = csv.reader(csvfile, delimiter=",")
    
        for word in csvreader:
            local_words.append(word)
    [unlimited_words] = local_words
    words = random.sample(unlimited_words, WORD_LIMIT)

def init_board():
    for y in range(INITIAL_HEIGHT):
        board.append([])
        for x in range(INITIAL_WIDTH):
            board[y].append("_")

def get_coords(word, off):
    global init
    if not init:
        size = len(word)
        x = int((INITIAL_WIDTH - size) / 2)
        y = int(INITIAL_HEIGHT / 2)
        init = True
        return [(x, y, False)]
    else:
        coords = []
        for y in range(INITIAL_HEIGHT):
            for x in range(INITIAL_WIDTH):
                if board[x][y] == word[off]:
                    coords.append((x, y - off, True))
                    coords.append((x - off, y, False))
        return coords

def check_collision(word, coords):
    # check board bounding box (FIXME)
    if coords[2]:
        if coords[1] + len(word) > INITIAL_HEIGHT:
            return True
    else:
        if coords[0] + len(word) > INITIAL_WIDTH:
            return True 

    # check letter collisions
    x = coords[0]
    y = coords[1]
    dir = coords[2]
    if dir:
        for i in range(0, len(word)):
            if board[x][y+i] != "_" and board[x][y+i] != word[i]:
                return True
    else:
        for i in range(0, len(word)):
            if board[x+i][y] != "_" and board[x+i][y] != word[i]:
                return True
    
    return False

def place_word(word, coords):
    x = coords[0]
    y = coords[1]
    dir = coords[2]
    if dir:
        for i in range(0, len(word)):
            board[x][y+i] = word[i]
    else:
        for i in range(0, len(word)):
            board[x+i][y] = word[i]

def check_word(word):
    for off in range(len(word)):
        coords = get_coords(word, off)
        if len(coords) == 0:
            # no common letters, skip by default (TODO: place where no collisions)
            continue
        for coord in coords:
            collided = check_collision(word, coord)
            if not collided:
                place_word(word, coord)
                return

def print_board():
    for y in range(len(board)):
        for x in range(len(board[0])):
            print(str(board[y][x]) + " ", end='')
        print("")
    print("")

def adjust_board():
    global board
    board_trimmed = []
    maxY, minY, maxX, minX = 0, INITIAL_HEIGHT, 0, INITIAL_WIDTH
    for y in range(INITIAL_HEIGHT):
        for x in range(INITIAL_WIDTH): 
            if board[y][x] != '_':
                if maxY < y:
                    maxY = y
                if minY > y:
                    minY = y
                if maxX < x:
                    maxX = x
                if minX > x:
                    minX = x

    i = 0
    for y in range(minY, maxY + 1):
        board_trimmed.append([])
        for x in range(minX, maxX + 1):
            if (board[y][x] == '_'):
                board_trimmed[i].append(random.choice('AĄBCĆDEĘFGHIJKLŁMNŃOÓPQRSŚTUVWXYZŹŻ'))
            else:
                board_trimmed[i].append(board[y][x].upper())
        i += 1

    board.clear()
    board = board_trimmed.copy() 

def create_doc():
    global board
    boardDataFrame = pd.DataFrame(board)
    currentTime = datetime.now()
    currentTimeString = currentTime.strftime("%Y_%m_%d_%H_%M_%S")
    
    doc = docx.Document()
    doc.add_paragraph(f"Odnajdź w gąszczu literek i zakreśl następujące wyrazy: {words}")
    table = doc.add_table(rows=boardDataFrame.shape[0], cols=boardDataFrame.shape[1])
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i in range(boardDataFrame.shape[0]):
        for j in range(boardDataFrame.shape[1]):
            cell = boardDataFrame.iat[i, j]
            table.cell(i, j).text = str(cell)
            table.cell(i, j).width = 1
            table.cell(i, j).height = 1

    doc.add_paragraph(f"W nagrodę możesz pokolorwać poniższy obrazek.")
    last_paragraph = doc.paragraphs[-1] 
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    lista_kolorowanek = os.listdir('./kolorowanki')
    kolorowanka = random.choice(lista_kolorowanek)

    image_path=f"./kolorowanki/{kolorowanka}"
    doc.add_picture(image_path, width=Inches(3.0), height=Inches(3.0))
    last_paragraph = doc.paragraphs[-1] 
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(f"./skreslanka-{currentTimeString }.docx")
    print(f"skreślankę zapisano do pliku ./skreslanka-{currentTimeString }.docx")

def generate_board():
    global init, words, board
    init_board()
    import_words()
    for word in words:
        check_word(word)
    adjust_board()
    if len(board) > HEIGHT_LIMIT or len(board[0]) > WIDTH_LIMIT:
        print(f"tabela wyszła zbyt duża (y={len(board)} * x={len(board[0])} )... ponawiam generowanie.")
        words.clear()
        board.clear()
        init = False
        init_board()
        generate_board()
    if len(board) - 2 > len(board[0]) or len(board[0]) - 2 > len(board):
        print(f"kształt tabeli za bardzo odbiega od kwadratu (y={len(board)} * x={len(board[0])} )... ponawiam generowanie.")
        words.clear()
        board.clear()
        init = False
        init_board()
        generate_board()

generate_board()
print_board()
create_doc()